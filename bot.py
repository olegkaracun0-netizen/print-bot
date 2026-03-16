#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Print Bot v6.0
- Смешанные заказы (фото + документы)
- Минимум 20 штук (общее)
- Копии отдельно для фото и для документов
- Чистый минималистичный дизайн
"""

import os, sys, logging, tempfile, json, re, shutil
import traceback, zipfile, threading, urllib.request, urllib.error
import time as _time
from datetime import datetime
from flask import (Flask, request, jsonify, send_file,
                   send_from_directory, render_template_string, abort)
import telegram
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, InputMediaPhoto
from telegram.ext import (Updater, CommandHandler, MessageHandler,
                           CallbackQueryHandler, ConversationHandler, Filters)
import PyPDF2
from docx import Document

# ═══════════════════════════════════════════
#  НАСТРОЙКИ
# ═══════════════════════════════════════════
TOKEN              = os.environ.get("TOKEN")
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
RENDER_URL         = os.environ.get("RENDER_EXTERNAL_URL")
PORT               = int(os.environ.get("PORT", 10000))
ADMIN_CHAT_ID      = 483613049
CONTACT_PHONE      = "89219805705"
MIN_ITEMS          = 20   # минимальный заказ

if not TOKEN:      print("❌ TOKEN не задан!");             sys.exit(1)
if not RENDER_URL: print("❌ RENDER_EXTERNAL_URL не задан!"); sys.exit(1)

ORDERS_PATH     = os.path.join(os.path.dirname(os.path.abspath(__file__)), "заказы")
ORDERS_DB_FILE  = os.path.join(ORDERS_PATH, "orders_history.json")
CLIENTS_DB_FILE = os.path.join(ORDERS_PATH, "clients.json")
os.makedirs(ORDERS_PATH, exist_ok=True)

ORDER_STATUSES = {
    "new":        "🆕 Новый",
    "processing": "🔄 В обработке",
    "printing":   "🖨 В печати",
    "ready":      "✅ Готов",
    "shipped":    "📦 Отправлен",
    "delivered":  "🏁 Доставлен",
    "cancelled":  "❌ Отменён",
}

DELIVERY_METHODS = {
    "pickup": ("🏪 Самовывоз СПб",   "Бесплатно"),
    "cdek":   ("📦 СДЭК",            "По тарифу"),
    "yandex": ("🚕 Яндекс Доставка", "По тарифу"),
}

LOYALTY_DISCOUNTS = {5: 5, 10: 10, 20: 15}

# ═══════════════════════════════════════════
#  ПРАЙС
# ═══════════════════════════════════════════
PHOTO_PRICES = {
    "small":  {(1,8):30,   (9,49):30,   (50,99):25,  (100,999999):15},
    "medium": {(1,8):60,   (9,49):60,   (50,99):50,  (100,999999):27},
    "large":  {(1,3):170,  (4,19):170,  (20,49):150, (50,999999):110},
}
DOC_PRICES = {
    "bw":    {(1,19):20,  (20,49):20,  (50,99):15,  (100,299):10, (300,999999):8},
    "color": {(1,19):40,  (20,49):40,  (50,99):30,  (100,299):20, (300,999999):16},
}

PRICE_INFO = """📋 *Прайс-лист*

📄 *Документы А4 — Чёрно-белая:*
  20–49 листов  →  20 ₽/лист
  50–99          →  15 ₽/лист
  100–299        →  10 ₽/лист
  от 300         →  8 ₽/лист

📄 *Документы А4 — Цветная:*
  20–49 листов  →  40 ₽/лист
  50–99          →  30 ₽/лист
  100–299        →  20 ₽/лист
  от 300         →  16 ₽/лист

📸 *Фото 10×15 / A6:*
  9–49 шт   →  30 ₽/шт
  50–99      →  25 ₽/шт
  100+        →  15 ₽/шт

📸 *Фото 13×18 / 15×21:*
  9–49 шт   →  60 ₽/шт
  50–99      →  50 ₽/шт
  100+        →  27 ₽/шт

📸 *Фото A4 / 21×30:*
  4–19 шт   →  170 ₽/шт
  20–49      →  150 ₽/шт
  50+         →  110 ₽/шт

📂 *Копии и сканы* — от 3 ₽/шт
💛 Большой заказ — обсудим условия

⚡️ *Минимальный заказ:* 20 шт/листов"""

AI_SYSTEM_PROMPT = f"""Ты вежливый ИИ-ассистент сервиса быстрой печати фото и документов.
Отвечай чётко, кратко, на русском языке, с эмодзи.
Телефон: {CONTACT_PHONE}. Доставка: Самовывоз СПб (бесплатно), СДЭК, Яндекс.
Форматы файлов: JPG, PNG, PDF, DOC, DOCX.
Минимальный заказ: {MIN_ITEMS} шт. Срок: 1–3 дня.
Скидки: от 5 заказов — 5%, от 10 — 10%, от 20 — 15%.
В один заказ можно добавить и фото и документы — бот посчитает всё по своим прайсам.
Документы Ч/Б: от 20л—20₽, от 50—15₽, от 100—10₽, от 300—8₽.
Документы Цвет: от 20л—40₽, от 50—30₽, от 100—20₽, от 300—16₽.
Фото 10×15: 9–49шт—30₽, 50–99—25₽, 100+—15₽.
Фото 13×18: 9–49шт—60₽, 50–99—50₽, 100+—27₽.
Фото A4: 4–19шт—170₽, 20–49—150₽, 50+—110₽.
Никогда не выдумывай цены. При вопросах не по теме — переводи на услуги."""

# ═══════════════════════════════════════════
#  ЛОГИРОВАНИЕ
# ═══════════════════════════════════════════
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s %(levelname)s %(message)s",
                    stream=sys.stdout)
logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════
#  СОСТОЯНИЯ ConversationHandler
# ═══════════════════════════════════════════
(
    COLLECTING_FILES,    # загрузка файлов
    SEL_PHOTO_FORMAT,    # выбор формата фото
    SEL_PHOTO_QTY,       # копии для фото
    SEL_DOC_TYPE,        # тип документов
    SEL_DOC_QTY,         # копии для документов
    SEL_DELIVERY,        # способ доставки
    ENTER_ADDRESS,       # ввод адреса
    CONFIRMING,          # подтверждение
    AI_CHAT,             # чат с ИИ
) = range(9)

# ═══════════════════════════════════════════
#  ГЛОБАЛЬНЫЕ ОБЪЕКТЫ
# ═══════════════════════════════════════════
user_sessions = {}
ai_histories  = {}
media_groups  = {}
group_timers  = {}
bot = updater = dispatcher = None

# ═══════════════════════════════════════════
#  УТИЛИТЫ
# ═══════════════════════════════════════════
def get_status_display(s):  return ORDER_STATUSES.get(s, s)
def extract_number(t):
    n = re.findall(r"\d+", t)
    return int(n[0]) if n else None
def format_size(b):
    if b < 1024:    return f"{b} B"
    if b < 1048576: return f"{b/1024:.1f} KB"
    return f"{b/1048576:.1f} MB"
def estimate_delivery(n):
    if n <= 50:  return "1 день"
    if n <= 200: return "2 дня"
    return "3 дня"

def calculate_price(pd, qty):
    for (mn, mx), p in pd.items():
        if mn <= qty <= mx:
            return p * qty
    return 0

def unit_price(pd, qty):
    for (mn, mx), p in pd.items():
        if mn <= qty <= mx:
            return p
    return 0

def next_tier(pd, qty):
    """Возвращает (следующий порог, цена на нём) или None."""
    tiers = sorted(pd.items(), key=lambda x: x[0][0])
    for (mn, mx), p in tiers:
        if mn > qty:
            return mn, p
    return None, None

def line():  return "─" * 22
def dot():   return "· · · · · · · · · · · · ·"

def count_items(path, name):
    try:
        if name.lower().endswith(".pdf"):
            with open(path, "rb") as f:
                return len(PyPDF2.PdfReader(f).pages), "страниц", "документ"
        if name.lower().endswith((".docx", ".doc")):
            d = Document(path)
            p = max(1, len(d.paragraphs) // 35) + len(d.tables) // 2
            return p, "страниц", "документ"
        if name.lower().endswith((".jpg", ".jpeg", ".png")):
            return 1, "фото", "фото"
    except Exception as e:
        logger.error(f"count_items: {e}")
    return 1, "единиц", "файл"

def download_file(fobj, fname):
    try:
        tmp  = tempfile.mkdtemp()
        path = os.path.join(tmp, fname)
        if hasattr(fobj, "get_file"):   fobj.get_file().download(custom_path=path)
        elif hasattr(fobj, "download"): fobj.download(custom_path=path)
        else:
            with open(path, "wb") as f: f.write(fobj.download_as_bytearray())
        return path, tmp
    except Exception as e:
        logger.error(f"download: {e}")
        return None, None

# ═══════════════════════════════════════════
#  КЛИЕНТСКАЯ БАЗА
# ═══════════════════════════════════════════
def load_clients():
    try:
        if os.path.exists(CLIENTS_DB_FILE):
            with open(CLIENTS_DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except: pass
    return {}

def save_clients(c):
    try:
        with open(CLIENTS_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(c, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"save_clients: {e}")

def get_client(uid):
    return load_clients().get(str(uid))

def upsert_client(uid, info, order=None):
    clients = load_clients()
    k = str(uid)
    if k not in clients:
        clients[k] = {"user_id": uid, "username": "", "first_name": "",
                      "orders": [], "total_spent": 0,
                      "joined": datetime.now().isoformat(),
                      "last_address": "", "last_delivery": "pickup"}
    clients[k].update({"username":   info.get("username", ""),
                        "first_name": info.get("first_name", ""),
                        "last_seen":  datetime.now().isoformat()})
    if order:
        clients[k]["orders"].append(order)
        clients[k]["total_spent"] += order.get("total", 0)
        if order.get("address"):         clients[k]["last_address"]  = order["address"]
        if order.get("delivery_method"): clients[k]["last_delivery"] = order["delivery_method"]
    save_clients(clients)
    return clients[k]

def loyalty_discount(uid):
    c = get_client(uid)
    if not c: return 0
    n = len(c.get("orders", []))
    d = 0
    for thr, pct in sorted(LOYALTY_DISCOUNTS.items()):
        if n >= thr: d = pct
    return d

def last_order(uid):
    c = get_client(uid)
    if c and c.get("orders"): return c["orders"][-1]
    return None

# ═══════════════════════════════════════════
#  СЕССИИ
# ═══════════════════════════════════════════
def init_session(uid, user):
    user_sessions[uid] = {
        "files":        [],    # все файлы
        "temp_dirs":    [],
        "photos":       [],    # только фото-файлы
        "docs":         [],    # только документы
        "total_photos": 0,     # сумма фото-штук
        "total_pages":  0,     # сумма страниц
        # настройки фото
        "photo_format": None,
        "photo_qty":    1,
        # настройки документов
        "doc_color":    None,
        "doc_qty":      1,
        # итог
        "total":        0,
        "discount":     0,
        "delivery":     "",
        "delivery_method": "pickup",
        "address":      "",
        "user_info": {
            "user_id":    uid,
            "username":   user.username or user.first_name,
            "first_name": user.first_name,
            "last_name":  user.last_name or "",
        },
    }
    upsert_client(uid, {"username":   user.username or user.first_name,
                         "first_name": user.first_name})

def cleanup(uid):
    if uid in user_sessions:
        for d in user_sessions[uid].get("temp_dirs", []):
            shutil.rmtree(d, ignore_errors=True)
        del user_sessions[uid]

# ═══════════════════════════════════════════
#  ИИ-АССИСТЕНТ (OpenRouter)
# ═══════════════════════════════════════════
def ask_ai(uid, msg):
    if not OPENROUTER_API_KEY:
        return f"🤖 ИИ не настроен. По вопросам: {CONTACT_PHONE}"
    if uid not in ai_histories: ai_histories[uid] = []
    ai_histories[uid].append({"role": "user", "content": msg})
    if len(ai_histories[uid]) > 20:
        ai_histories[uid] = ai_histories[uid][-20:]
    msgs = [{"role": "system", "content": AI_SYSTEM_PROMPT}] + ai_histories[uid]
    try:
        payload = json.dumps({"model": "openrouter/free", "messages": msgs,
                              "max_tokens": 800, "temperature": 0.7}).encode()
        req = urllib.request.Request(
            "https://openrouter.ai/api/v1/chat/completions", data=payload,
            headers={"Content-Type": "application/json",
                     "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                     "HTTP-Referer": RENDER_URL, "X-Title": "Print Bot"},
            method="POST")
        with urllib.request.urlopen(req, timeout=30) as r:
            data = json.loads(r.read())
        reply = data["choices"][0]["message"]["content"]
        ai_histories[uid].append({"role": "assistant", "content": reply})
        return reply
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        logger.error(f"AI {e.code}: {body[:200]}")
        return f"😔 Ошибка ИИ. По вопросам: {CONTACT_PHONE}"
    except Exception as e:
        logger.error(f"AI: {e}")
        return f"😔 Нет ответа от ИИ. По вопросам: {CONTACT_PHONE}"

# ═══════════════════════════════════════════
#  ИСТОРИЯ ЗАКАЗОВ
# ═══════════════════════════════════════════
def load_history():
    try:
        if os.path.exists(ORDERS_DB_FILE):
            with open(ORDERS_DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except: pass
    return []

def save_history(entry):
    try:
        h = load_history(); h.append(entry)
        with open(ORDERS_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(h, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logger.error(f"save_history: {e}"); return False

def update_order_status(oid, status):
    try:
        h = load_history(); uid = None
        for o in h:
            if o.get("order_id") == oid:
                o["status"] = status; uid = o.get("user_id"); break
        with open(ORDERS_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(h, f, ensure_ascii=False, indent=2)
        inf = os.path.join(ORDERS_PATH, oid, "информация_о_заказе.txt")
        if os.path.exists(inf):
            with open(inf, "r", encoding="utf-8") as f: txt = f.read()
            txt = re.sub(r"Статус:.*\n", f"Статус: {get_status_display(status)}\n", txt)
            with open(inf, "w", encoding="utf-8") as f: f.write(txt)
        if uid and bot:
            bot.send_message(chat_id=uid,
                text=f"🔔 *Статус заказа обновлён*\n\n🆔 `{oid}`\n📌 *{get_status_display(status)}*",
                parse_mode="Markdown")
        return True
    except Exception as e:
        logger.error(f"update_status: {e}"); return False

# ═══════════════════════════════════════════
#  СОХРАНЕНИЕ ЗАКАЗА
# ═══════════════════════════════════════════
def save_order(uid, s):
    try:
        uname  = re.sub(r"[^\w\s-]", "", s["user_info"]["username"]) or f"user_{uid}"
        ts     = datetime.now().strftime("%Y%m%d_%H%M%S")
        oid    = f"{uname}_{ts}"
        folder = os.path.join(ORDERS_PATH, oid)
        os.makedirs(folder, exist_ok=True)

        for i, f in enumerate(s["files"], 1):
            if os.path.exists(f["path"]):
                safe = re.sub(r'[<>:"/\\|?*]', "", f["name"])
                shutil.copy2(f["path"], os.path.join(folder, f"{i}_{safe}"))

        ph = s["photos"]; dc = s["docs"]
        tp = s["total_photos"]; td = s["total_pages"]
        dlabel = DELIVERY_METHODS.get(s["delivery_method"], ("Самовывоз",""))[0]

        with open(os.path.join(folder, "информация_о_заказе.txt"), "w", encoding="utf-8") as f:
            f.write("ЗАКАЗ " + datetime.now().strftime("%d.%m.%Y %H:%M:%S") + "\n" + "="*50 + "\n\n")
            f.write(f"Клиент: {s['user_info']['first_name']} (@{uname})\nID: {uid}\n\n")
            if ph:
                fn = {"small":"10×15/A6","medium":"13×18","large":"A4"}
                f.write(f"ФОТО: {len(ph)} файлов, {tp} шт.\n")
                f.write(f"  Формат: {fn.get(s['photo_format'],'')}\n")
                f.write(f"  Копий: {s['photo_qty']}\n")
                f.write(f"  К печати: {tp * s['photo_qty']} шт.\n\n")
            if dc:
                cn = {"bw":"Ч/Б","color":"Цветная"}
                f.write(f"ДОКУМЕНТЫ: {len(dc)} файлов, {td} стр.\n")
                f.write(f"  Тип: {cn.get(s['doc_color'],'')}\n")
                f.write(f"  Копий: {s['doc_qty']}\n")
                f.write(f"  К печати: {td * s['doc_qty']} стр.\n\n")
            if s.get("discount"):
                f.write(f"Скидка: {s['discount']}%\n")
            f.write(f"ИТОГО: {s['total']} руб.\n")
            f.write(f"Доставка: {dlabel}\n")
            if s.get("address"): f.write(f"Адрес: {s['address']}\n")
            f.write(f"Статус: {get_status_display('new')}\n")

        save_history({
            "order_id": oid, "folder": folder,
            "user_id": uid, "username": uname,
            "user_name": s["user_info"]["first_name"],
            "date": datetime.now().isoformat(),
            "has_photos": bool(ph), "has_docs": bool(dc),
            "total_photos": tp, "total_pages": td,
            "photo_qty": s["photo_qty"], "doc_qty": s["doc_qty"],
            "total_price": s["total"], "discount": s.get("discount", 0),
            "delivery": estimate_delivery(tp + td),
            "delivery_method": s["delivery_method"],
            "address": s.get("address", ""),
            "status": "new",
        })
        upsert_client(uid, s["user_info"], {
            "order_id": oid, "total": s["total"],
            "delivery_method": s["delivery_method"],
            "address": s.get("address", ""),
        })
        return True, oid
    except Exception as e:
        logger.error(f"save_order: {e}\n{traceback.format_exc()}")
        return False, None

# ═══════════════════════════════════════════
#  УВЕДОМЛЕНИЕ АДМИНУ
# ═══════════════════════════════════════════
def notify_admin(s, oid):
    try:
        url    = f"{RENDER_URL}/orders/{oid}/"
        ph     = s["photos"]; dc = s["docs"]
        tp     = s["total_photos"]; td = s["total_pages"]
        dlabel = DELIVERY_METHODS.get(s["delivery_method"], ("Самовывоз",""))[0]
        fn     = {"small":"10×15/A6","medium":"13×18","large":"A4"}
        cn     = {"bw":"Ч/Б","color":"Цветная"}

        lines = [
            "🔔 *НОВЫЙ ЗАКАЗ*", "",
            f"👤 *{s['user_info']['first_name']}*  @{s['user_info']['username']}",
            f"🆔 `{s['user_info']['user_id']}`", "",
        ]
        if ph:
            lines.append(f"📸 Фото {fn.get(s['photo_format'],'')} — {tp} шт. × {s['photo_qty']} коп.")
        if dc:
            lines.append(f"📄 Документы {cn.get(s['doc_color'],'')} — {td} стр. × {s['doc_qty']} коп.")
        if s.get("discount"):
            lines.append(f"🎁 Скидка: *{s['discount']}%*")
        lines += [
            "",
            f"💰 *{s['total']} руб.*",
            f"🚚 {dlabel}",
        ]
        if s.get("address"):
            lines.append(f"📍 {s['address']}")
        lines += ["", f"🔗 [Открыть заказ]({url})"]
        if bot:
            bot.send_message(chat_id=ADMIN_CHAT_ID,
                             text="\n".join(lines), parse_mode="Markdown")
    except Exception as e:
        logger.error(f"notify_admin: {e}")

# ═══════════════════════════════════════════
#  КЛАВИАТУРЫ  (новый дизайн)
# ═══════════════════════════════════════════
def kbd_main(uid=None):
    has_last = bool(last_order(uid)) if uid else False
    rows = [
        [InlineKeyboardButton("📎  Отправить файлы на печать", callback_data="start_order")],
        [InlineKeyboardButton("💰  Цены",           callback_data="show_prices"),
         InlineKeyboardButton("📋  Мои заказы",     callback_data="my_orders")],
        [InlineKeyboardButton("🤖  Помощник ИИ",    callback_data="start_ai"),
         InlineKeyboardButton("📞  Контакты",        callback_data="show_contacts")],
    ]
    if has_last:
        rows.insert(1, [InlineKeyboardButton("🔄  Повторить последний заказ", callback_data="repeat_order")])
    return InlineKeyboardMarkup(rows)

def kbd_photo_format():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("◾ 10×15 / A6  —  малый",     callback_data="pf_small")],
        [InlineKeyboardButton("◽ 13×18 / 15×21  —  средний", callback_data="pf_medium")],
        [InlineKeyboardButton("⬜ A4 / 21×30  —  большой",   callback_data="pf_large")],
        [InlineKeyboardButton("➕  Добавить ещё файлы",       callback_data="add_more")],
        [InlineKeyboardButton("✖  Отменить заказ",            callback_data="cancel")],
    ])

def kbd_doc_type():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("◼  Чёрно-белая печать",   callback_data="dt_bw")],
        [InlineKeyboardButton("🎨  Цветная печать",       callback_data="dt_color")],
        [InlineKeyboardButton("➕  Добавить ещё файлы",   callback_data="add_more")],
        [InlineKeyboardButton("✖  Отменить заказ",        callback_data="cancel")],
    ])

def kbd_qty(prefix):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("1",  callback_data=f"{prefix}_1"),
         InlineKeyboardButton("2",  callback_data=f"{prefix}_2"),
         InlineKeyboardButton("3",  callback_data=f"{prefix}_3"),
         InlineKeyboardButton("5",  callback_data=f"{prefix}_5"),
         InlineKeyboardButton("10", callback_data=f"{prefix}_10")],
        [InlineKeyboardButton("20", callback_data=f"{prefix}_20"),
         InlineKeyboardButton("30", callback_data=f"{prefix}_30"),
         InlineKeyboardButton("50", callback_data=f"{prefix}_50"),
         InlineKeyboardButton("100",callback_data=f"{prefix}_100")],
        [InlineKeyboardButton("✏️  Написать число в чат", callback_data=f"{prefix}_hint")],
        [InlineKeyboardButton("✖  Отменить заказ",         callback_data="cancel")],
    ])

def kbd_delivery(uid=None):
    c = get_client(uid) if uid else None
    last = c.get("last_delivery", "pickup") if c else "pickup"
    rows = []
    for k, (label, price) in DELIVERY_METHODS.items():
        mark = "  ✓" if k == last else ""
        rows.append([InlineKeyboardButton(f"{label}  —  {price}{mark}",
                                          callback_data=f"dlv_{k}")])
    rows.append([InlineKeyboardButton("✖  Отменить", callback_data="cancel")])
    return InlineKeyboardMarkup(rows)

def kbd_confirm():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✅  Подтвердить заказ",  callback_data="confirm")],
        [InlineKeyboardButton("✖  Отменить",            callback_data="cancel")],
    ])

def kbd_after_order(uid=None):
    rows = [
        [InlineKeyboardButton("📎  Новый заказ",       callback_data="new_order")],
        [InlineKeyboardButton("🤖  Помощник ИИ",       callback_data="start_ai")],
    ]
    if uid and last_order(uid):
        rows.insert(1, [InlineKeyboardButton("🔄  Повторить заказ", callback_data="repeat_order")])
    return InlineKeyboardMarkup(rows)

def kbd_ai():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📎  Сделать заказ",  callback_data="start_order"),
         InlineKeyboardButton("💰  Цены",           callback_data="show_prices")],
        [InlineKeyboardButton("🔄  Сбросить диалог", callback_data="reset_ai"),
         InlineKeyboardButton("🏠  Меню",            callback_data="main_menu")],
    ])

# ═══════════════════════════════════════════
#  ОБРАБОТКА ФАЙЛОВ
# ═══════════════════════════════════════════
def handle_file(update, context):
    if update.message.media_group_id:
        return handle_media_group(update, context)
    return process_single_file(update, context)

def handle_media_group(update, context):
    uid  = update.effective_user.id
    mgid = update.message.media_group_id
    media_groups.setdefault(uid, {}).setdefault(mgid, []).append(update.message)
    key = f"{uid}_{mgid}"
    if key in group_timers: group_timers[key].cancel()
    t = threading.Timer(2.0, flush_group, args=[uid, mgid, context])
    t.daemon = True; t.start()
    group_timers[key] = t
    return COLLECTING_FILES

def flush_group(uid, mgid, context):
    try:
        if uid not in media_groups or mgid not in media_groups[uid]: return
        msgs = media_groups[uid].pop(mgid)
        group_timers.pop(f"{uid}_{mgid}", None)
        if uid not in user_sessions: init_session(uid, msgs[0].from_user)
        for msg in msgs: _add_file_from_msg(uid, msg)
        # state от flush_group нельзя вернуть напрямую (вызывается из Timer),
        # поэтому сохраняем в сессии
        next_st = send_file_summary(uid, context)
        if next_st and uid in user_sessions:
            user_sessions[uid]["_conv_state"] = next_st
    except Exception as e:
        logger.error(f"flush_group: {e}\n{traceback.format_exc()}")

def process_single_file(update, context):
    uid = update.effective_user.id
    msg = update.message
    if uid not in user_sessions: init_session(uid, update.effective_user)
    _add_file_from_msg(uid, msg)
    return send_file_summary(uid, context, reply_fn=msg.reply_text)

def _add_file_from_msg(uid, msg):
    s = user_sessions.get(uid)
    if not s: return
    fobj = fname = ftype = None
    if msg.document:
        fobj  = msg.document; fname = fobj.file_name
        ext   = fname.lower().split(".")[-1]
        if ext in ("jpg","jpeg","png"): ftype = "photo"
        elif ext in ("pdf","doc","docx"): ftype = "doc"
        else: return
    elif msg.photo:
        fobj  = msg.photo[-1]
        fname = f"photo_{datetime.now().strftime('%Y%m%d_%H%M%S%f')}.jpg"
        ftype = "photo"
    else: return

    path, tmp = download_file(fobj, fname)
    if not path: return
    items, unit, _ = count_items(path, fname)
    entry = {"path": path, "name": fname, "type": ftype, "items": items, "unit": unit}
    s["files"].append(entry)
    s["temp_dirs"].append(tmp)
    if ftype == "photo":
        s["photos"].append(entry)
        s["total_photos"] += items
    else:
        s["docs"].append(entry)
        s["total_pages"] += items

def send_file_summary(uid, context, reply_fn=None):
    """Показывает итог загрузки + прайс + следующий шаг."""
    s  = user_sessions.get(uid)
    if not s: return
    ph = s["photos"]; dc = s["docs"]
    tp = s["total_photos"]; td = s["total_pages"]
    total_items = tp + td

    lines = ["🎉 *Файлы приняты!*", ""]
    if ph: lines.append(f"📸 Фото:       {len(ph)} файл(а)  /  {tp} шт.")
    if dc: lines.append(f"📄 Документы:  {len(dc)} файл(а)  /  {td} стр.")
    lines.append(f"📦 Всего:      {total_items} шт.")
    lines.append("")
    lines.append(dot())
    lines.append("")

    if total_items < MIN_ITEMS:
        need = MIN_ITEMS - total_items
        lines += [
            f"⚠️ *Минимальный заказ — {MIN_ITEMS} шт.*",
            f"➕ Добавь ещё *{need} шт.* — и можно заказывать!",
            "",
        ]
        kbd = InlineKeyboardMarkup([
            [InlineKeyboardButton("➕  Добавить файлы",  callback_data="add_more")],
            [InlineKeyboardButton("✖  Отменить",         callback_data="cancel")],
        ])
        txt = "\n".join(lines)
        if reply_fn: reply_fn(txt, reply_markup=kbd, parse_mode="Markdown")
        else: context.bot.send_message(chat_id=uid, text=txt, reply_markup=kbd, parse_mode="Markdown")
        return COLLECTING_FILES

    # Показываем актуальный прайс
    if ph:
        lines.append("📸 *Цены на фото* (за штуку):")
        lines.append("  10×15  →  30₽ / 25₽ / 15₽")
        lines.append("  13×18  →  60₽ / 50₽ / 27₽")
        lines.append("  A4     →  170₽ / 150₽ / 110₽")
        lines.append("")
    if dc:
        lines.append("📄 *Цены на документы* (за лист):")
        lines.append("  Ч/Б    →  20₽ / 15₽ / 10₽ / 8₽")
        lines.append("  Цвет   →  40₽ / 30₽ / 20₽ / 16₽")
        lines.append("")

    lines.append(dot())
    lines.append("")

    # Следующий шаг
    if ph:
        lines.append("🖼 *Выбери формат для фото:*")
        kbd = kbd_photo_format()
        next_state = SEL_PHOTO_FORMAT
    else:
        lines.append("🖨 *Выбери тип печати документов:*")
        kbd = kbd_doc_type()
        next_state = SEL_DOC_TYPE

    txt = "\n".join(lines)
    if reply_fn: reply_fn(txt, reply_markup=kbd, parse_mode="Markdown")
    else: context.bot.send_message(chat_id=uid, text=txt, reply_markup=kbd, parse_mode="Markdown")
    return next_state

# ═══════════════════════════════════════════
#  РАСЧЁТ ИТОГА
# ═══════════════════════════════════════════
def calc_total(uid):
    s = user_sessions.get(uid)
    if not s: return 0
    total = 0
    ph = s["photos"]; dc = s["docs"]
    tp = s["total_photos"]; td = s["total_pages"]

    if ph and s.get("photo_format") and s.get("photo_qty"):
        pd     = PHOTO_PRICES[s["photo_format"]]
        qty    = s["photo_qty"]
        ftotal = sum(calculate_price(pd, f["items"] * qty) for f in ph)
        total += ftotal

    if dc and s.get("doc_color") and s.get("doc_qty"):
        pd     = DOC_PRICES[s["doc_color"]]
        qty    = s["doc_qty"]
        ftotal = calculate_price(pd, td * qty)
        total += ftotal

    disc     = loyalty_discount(uid)
    s["discount"] = disc
    if disc: total = int(total * (100 - disc) / 100)
    s["total"] = total
    return total

def build_order_summary(uid):
    """Строит красивый текст итога заказа."""
    s  = user_sessions.get(uid)
    if not s: return ""
    ph = s["photos"]; dc = s["docs"]
    tp = s["total_photos"]; td = s["total_pages"]
    fn = {"small":"10×15 / A6","medium":"13×18 / 15×21","large":"A4 / 21×30"}
    cn = {"bw":"Чёрно-белая","color":"Цветная"}
    dlabel = DELIVERY_METHODS.get(s["delivery_method"], ("Самовывоз",""))[0]

    lines = ["📋 *Ваш заказ*", ""]

    if ph:
        pq   = s["photo_qty"]
        pd   = PHOTO_PRICES[s["photo_format"]]
        phq  = tp * pq
        pp   = unit_price(pd, phq)
        ptot = pp * phq
        disc_ph = int(ptot * (100 - s["discount"]) / 100) if s["discount"] else ptot
        lines += [
            f"📸 *Фото  {fn.get(s['photo_format'],'')}*",
            f"   {tp} шт. × {pq} коп. = {phq} шт.",
            f"   {pp} ₽/шт  →  *{disc_ph} ₽*",
            "",
        ]

    if dc:
        dq   = s["doc_qty"]
        pd   = DOC_PRICES[s["doc_color"]]
        dcq  = td * dq
        dp   = unit_price(pd, dcq)
        dtot = dp * dcq
        disc_dc = int(dtot * (100 - s["discount"]) / 100) if s["discount"] else dtot
        lines += [
            f"📄 *Документы  {cn.get(s['doc_color'],'')}*",
            f"   {td} стр. × {dq} коп. = {dcq} стр.",
            f"   {dp} ₽/лист  →  *{disc_dc} ₽*",
            "",
        ]

    lines.append(dot())
    lines.append("")
    if s["discount"]:
        lines.append(f"👑 Скидка постоянного клиента:  *−{s['discount']}%*")
    lines.append(f"💰 *Итого:  {s['total']} ₽*")
    lines.append(f"⏱ Срок:    {estimate_delivery(tp + td)}")
    lines.append(f"🚚 Доставка: {dlabel}")
    if s.get("address"):
        lines.append(f"📍 {s['address']}")
    lines.append("")
    lines.append(dot())
    lines.append("")
    lines.append("Всё верно?")
    return "\n".join(lines)

# ═══════════════════════════════════════════
#  CANCEL
# ═══════════════════════════════════════════
def do_cancel(uid, query=None, context=None):
    cleanup(uid)
    text = "✖ *Заказ отменён*\n\nВсе файлы удалены."
    kbd  = kbd_after_order(uid)
    if query:
        try: query.edit_message_text(text, reply_markup=kbd, parse_mode="Markdown")
        except:
            if context: context.bot.send_message(chat_id=uid, text=text, reply_markup=kbd, parse_mode="Markdown")
    elif context:
        context.bot.send_message(chat_id=uid, text=text, reply_markup=kbd, parse_mode="Markdown")
    return COLLECTING_FILES

# ═══════════════════════════════════════════
#  КОМАНДЫ
# ═══════════════════════════════════════════
def cmd_start(update, context):
    uid  = update.effective_user.id
    user = update.effective_user
    cleanup(uid)
    c    = get_client(uid)
    n    = len(c.get("orders", [])) if c else 0
    disc = loyalty_discount(uid)

    if n > 0:
        hi = f"С возвращением, *{user.first_name}*! 👋\n🛍 Заказов: *{n}*"
    else:
        hi = f"Привет, *{user.first_name}*! 👋"
    if disc: hi += f"  |  👑 Скидка: *{disc}%*"

    text = (
        f"🖨✨ *Сервис быстрой печати*\n\n"
        f"{hi}\n\n"
        f"Что печатаем:\n"
        f"📸 Фото любых форматов\n"
        f"📄 Документы Ч/Б и цветные\n"
        f"📚 Курсовые, дипломы, рефераты\n"
        f"📂 Копии и сканирование\n\n"
        f"{dot()}\n\n"
        f"📎 Отправь файлы — и я всё посчитаю.\n"
        f"В один заказ можно и фото и документы.\n"
        f"*Минимум {MIN_ITEMS} штук/листов.*"
    )
    update.message.reply_text(text, reply_markup=kbd_main(uid), parse_mode="Markdown")
    return COLLECTING_FILES

def cmd_myorders(update, context):
    uid = update.effective_user.id
    c   = get_client(uid)
    if not c or not c.get("orders"):
        update.message.reply_text(
            "📋 *Мои заказы*\n\nПока заказов нет.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("📎  Сделать первый заказ", callback_data="start_order")]
            ])
        )
        return COLLECTING_FILES

    orders = c["orders"]; disc = loyalty_discount(uid)
    lines  = [
        "📋 *Мои заказы*", "",
        f"Всего:    *{len(orders)}*",
        f"Потрачено: *{c.get('total_spent',0)} ₽*",
    ]
    if disc: lines.append(f"👑 Скидка:  *{disc}%*")
    lines += ["", dot(), ""]
    for o in reversed(orders[-5:]):
        lines.append(f"🆔 `{o['order_id']}`  —  *{o.get('total',0)} ₽*")
        lines.append(f"   📅 {o['date'][:10]}")
        lines.append("")
    update.message.reply_text(
        "\n".join(lines), parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("🔄  Повторить последний", callback_data="repeat_order")],
            [InlineKeyboardButton("📎  Новый заказ",          callback_data="start_order")],
            [InlineKeyboardButton("🏠  Меню",                 callback_data="main_menu")],
        ])
    )
    return COLLECTING_FILES

def cmd_prices(update, context):
    update.message.reply_text(
        PRICE_INFO, parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("📎  Сделать заказ", callback_data="start_order")],
            [InlineKeyboardButton("🤖  Спросить ИИ",  callback_data="start_ai")],
        ])
    )
    return COLLECTING_FILES

def cmd_ai(update, context):
    update.message.reply_text(
        "🤖 *Помощник ИИ*\n\nЗадай любой вопрос о ценах, услугах и сроках 💬",
        parse_mode="Markdown", reply_markup=kbd_ai()
    )
    return AI_CHAT

def cmd_help(update, context):
    uid  = update.effective_user.id
    disc = loyalty_discount(uid)
    text = (
        f"📖 *Как сделать заказ:*\n\n"
        f"1. Отправь файлы (JPG, PNG, PDF, DOC, DOCX)\n"
        f"2. Выбери формат и количество копий\n"
        f"3. Выбери доставку\n"
        f"4. Подтверди заказ\n\n"
        f"{dot()}\n\n"
        f"*Команды:*\n"
        f"/start — Главное меню\n"
        f"/myorders — Мои заказы\n"
        f"/prices — Прайс\n"
        f"/ai — Помощник ИИ\n"
        f"/help — Справка\n\n"
        f"{dot()}\n\n"
        f"👑 *Скидки:*\n"
        f"от 5 заказов — 5%\n"
        f"от 10 заказов — 10%\n"
        f"от 20 заказов — 15%\n"
    )
    if disc: text += f"\nТвоя скидка сейчас: *{disc}%*"
    update.message.reply_text(text, parse_mode="Markdown", reply_markup=kbd_main(uid))
    return COLLECTING_FILES

# ═══════════════════════════════════════════
#  ИИ ЧАТ
# ═══════════════════════════════════════════
def ai_msg_handler(update, context):
    uid = update.effective_user.id
    context.bot.send_chat_action(chat_id=uid, action=telegram.ChatAction.TYPING)
    reply = ask_ai(uid, update.message.text)
    update.message.reply_text(
        f"🤖 {reply}\n\n{dot()}\n_Задай следующий вопрос или выбери действие:_",
        parse_mode="Markdown", reply_markup=kbd_ai()
    )
    return AI_CHAT

# ═══════════════════════════════════════════
#  ВВОД КОЛИЧЕСТВА (текстом)
# ═══════════════════════════════════════════
def qty_text_handler(update, context):
    """Обрабатывает ручной ввод числа. Определяет что сейчас вводится по сессии."""
    uid = update.effective_user.id
    n   = extract_number(update.message.text)
    s   = user_sessions.get(uid)
    if not s: return do_cancel(uid, context=context)

    # Определяем текущий шаг
    step = s.get("_qty_step", "photo")

    if not n or n < 1 or n > 10000:
        update.message.reply_text(
            "Введи число от *1 до 10000*",
            parse_mode="Markdown",
            reply_markup=kbd_qty("pq" if step == "photo" else "dq")
        )
        return SEL_PHOTO_QTY if step == "photo" else SEL_DOC_QTY

    return _process_qty(uid, n, step, context)

def _process_qty(uid, qty, step, context, query=None):
    s = user_sessions.get(uid)
    if not s: return do_cancel(uid, context=context)

    if step == "photo":
        s["photo_qty"] = qty
        # Если есть документы — переходим к их настройке
        if s["docs"]:
            s["_qty_step"] = "doc"
            cn = {"bw":"Чёрно-белая","color":"Цветная"}
            text = (
                f"📸 Фото: *{qty} коп.* — принято\n\n"
                f"{dot()}\n\n"
                f"📄 *Теперь настроим документы*\n\n"
                f"Выбери тип печати:"
            )
            kbd = kbd_doc_type()
            if query:
                try: query.edit_message_text(text, reply_markup=kbd, parse_mode="Markdown")
                except: context.bot.send_message(chat_id=uid, text=text, reply_markup=kbd, parse_mode="Markdown")
            else:
                context.bot.send_message(chat_id=uid, text=text, reply_markup=kbd, parse_mode="Markdown")
            return SEL_DOC_TYPE
        else:
            return _go_to_delivery(uid, context, query)

    elif step == "doc":
        s["doc_qty"] = qty
        return _go_to_delivery(uid, context, query)

def _go_to_delivery(uid, context, query=None):
    s = user_sessions.get(uid)
    calc_total(uid)
    text = (
        f"📦 *Способ доставки:*\n\n"
        f"Итого заказа: *{s['total']} ₽*"
    )
    kbd = kbd_delivery(uid)
    if query:
        try: query.edit_message_text(text, reply_markup=kbd, parse_mode="Markdown")
        except: context.bot.send_message(chat_id=uid, text=text, reply_markup=kbd, parse_mode="Markdown")
    else:
        context.bot.send_message(chat_id=uid, text=text, reply_markup=kbd, parse_mode="Markdown")
    return SEL_DELIVERY

# ═══════════════════════════════════════════
#  ВВОД АДРЕСА (текстом)
# ═══════════════════════════════════════════
def address_handler(update, context):
    uid = update.effective_user.id
    s   = user_sessions.get(uid)
    if not s: return do_cancel(uid, context=context)
    s["address"] = update.message.text.strip()
    text = build_order_summary(uid)
    update.message.reply_text(text, reply_markup=kbd_confirm(), parse_mode="Markdown")
    return CONFIRMING

# ═══════════════════════════════════════════
#  BUTTON HANDLER
# ═══════════════════════════════════════════
def button_handler(update, context):
    q    = update.callback_query
    q.answer()
    uid  = q.from_user.id
    data = q.data
    s    = user_sessions.get(uid)
    logger.info(f"CB {data} uid={uid}")

    # ── меню ────────────────────────────────
    if data == "main_menu":
        q.edit_message_text(
            "🖨 *Сервис печати*\n\nВыбери действие:",
            reply_markup=kbd_main(uid), parse_mode="Markdown"
        )
        return COLLECTING_FILES

    if data == "start_order":
        q.edit_message_text(
            f"📎 *Отправь файлы для печати*\n\n"
            f"Форматы: JPG, PNG, PDF, DOC, DOCX\n"
            f"Можно сразу несколько — и фото и документы.\n\n"
            f"_Минимум {MIN_ITEMS} штук/листов в заказе._",
            parse_mode="Markdown"
        )
        return COLLECTING_FILES

    if data == "show_prices":
        try: q.edit_message_text(PRICE_INFO, parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("📎  Заказать",   callback_data="start_order")],
                    [InlineKeyboardButton("🏠  Меню",       callback_data="main_menu")],
                ]))
        except: context.bot.send_message(chat_id=uid, text=PRICE_INFO, parse_mode="Markdown")
        return COLLECTING_FILES

    if data == "show_contacts":
        q.edit_message_text(
            f"📞 *Контакты*\n\n"
            f"Телефон:  *{CONTACT_PHONE}*\n\n"
            f"Доставка:\n"
            f"  🏪 Самовывоз СПб — бесплатно\n"
            f"  📦 СДЭК — по тарифу\n"
            f"  🚕 Яндекс Доставка — по тарифу\n\n"
            f"Режим работы: пн–сб 9:00–20:00",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("📎  Заказать", callback_data="start_order")],
                [InlineKeyboardButton("🏠  Меню",     callback_data="main_menu")],
            ]),
            parse_mode="Markdown"
        )
        return COLLECTING_FILES

    if data == "my_orders":
        c = get_client(uid)
        if not c or not c.get("orders"):
            q.edit_message_text("📋 Заказов пока нет.",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("📎  Сделать заказ", callback_data="start_order")]
                ]))
            return COLLECTING_FILES
        orders = c["orders"]; disc = loyalty_discount(uid)
        lines  = ["📋 *Мои заказы*", "",
                  f"Всего: *{len(orders)}*   Потрачено: *{c.get('total_spent',0)} ₽*"]
        if disc: lines.append(f"👑 Скидка: *{disc}%*")
        lines += ["", dot(), ""]
        for o in reversed(orders[-5:]):
            lines.append(f"🆔 `{o['order_id']}`  —  *{o.get('total',0)} ₽*")
            lines.append(f"   📅 {o['date'][:10]}")
            lines.append("")
        try:
            q.edit_message_text("\n".join(lines), parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("🔄  Повторить", callback_data="repeat_order")],
                    [InlineKeyboardButton("📎  Новый",     callback_data="start_order")],
                    [InlineKeyboardButton("🏠  Меню",      callback_data="main_menu")],
                ]))
        except: pass
        return COLLECTING_FILES

    if data == "repeat_order":
        lo = last_order(uid)
        if not lo:
            q.answer("Нет предыдущих заказов", show_alert=True)
            return COLLECTING_FILES
        q.edit_message_text(
            f"🔄 *Повтор заказа*\n\n"
            f"Последний заказ: *{lo.get('total',0)} ₽*\n"
            f"📅 {lo['date'][:10]}\n\n"
            f"Отправь файлы заново — и я всё посчитаю.\n"
            f"_Настройки формата и количества можно изменить._",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("🏠  Меню", callback_data="main_menu")]
            ])
        )
        return COLLECTING_FILES

    # ── ИИ ──────────────────────────────────
    if data == "start_ai":
        try:
            q.edit_message_text(
                "🤖 *Помощник ИИ*\n\nЗадай вопрос о ценах, услугах или сроках 💬",
                reply_markup=kbd_ai(), parse_mode="Markdown"
            )
        except:
            context.bot.send_message(chat_id=uid,
                text="🤖 *Помощник ИИ*\n\nЗадай вопрос 💬",
                reply_markup=kbd_ai(), parse_mode="Markdown")
        return AI_CHAT

    if data == "reset_ai":
        ai_histories.pop(uid, None)
        q.edit_message_text(
            "🤖 Диалог сброшен. Задай новый вопрос 💬",
            reply_markup=kbd_ai(), parse_mode="Markdown"
        )
        return AI_CHAT

    # ── файлы / отмена ───────────────────────
    if data == "cancel":
        return do_cancel(uid, q, context)

    if data == "add_more":
        q.edit_message_text(
            "📎 *Добавь файлы*\n\n_JPG, PNG, PDF, DOC, DOCX_",
            parse_mode="Markdown"
        )
        return COLLECTING_FILES

    if data == "new_order":
        cleanup(uid)
        q.edit_message_text(
            f"📎 *Новый заказ*\n\nОтправь файлы для печати.\n"
            f"_Минимум {MIN_ITEMS} штук/листов._",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("🏠  Меню", callback_data="main_menu")]
            ])
        )
        return COLLECTING_FILES

    # ── формат фото ─────────────────────────
    if data.startswith("pf_"):
        if not s: return do_cancel(uid, q, context)
        fmt = data[3:]
        s["photo_format"] = fmt
        s["_qty_step"]    = "photo"
        fn  = {"small":"10×15 / A6","medium":"13×18 / 15×21","large":"A4 / 21×30"}
        tp  = s["total_photos"]

        # Подсказка по цене
        pd   = PHOTO_PRICES[fmt]
        up   = unit_price(pd, tp)
        nt, np_ = next_tier(pd, tp)
        hint = f"💡 {tp} фото → *{up} ₽/шт*"
        if nt:
            hint += f"   |   от {nt} шт → *{np_} ₽/шт*"

        q.edit_message_text(
            f"📸 *Формат: {fn.get(fmt,fmt)}*\n\n"
            f"{hint}\n\n"
            f"{dot()}\n\n"
            f"Сколько *копий каждого фото* напечатать?\n\n"
            f"_Напиши число или выбери кнопку:_",
            parse_mode="Markdown",
            reply_markup=kbd_qty("pq")
        )
        return SEL_PHOTO_QTY

    # ── тип документов ───────────────────────
    if data.startswith("dt_"):
        if not s: return do_cancel(uid, q, context)
        color = data[3:]
        s["doc_color"]  = color
        s["_qty_step"]  = "doc"
        cn  = {"bw":"Чёрно-белая","color":"Цветная"}
        td  = s["total_pages"]
        pd  = DOC_PRICES[color]
        up  = unit_price(pd, td)
        nt, np_ = next_tier(pd, td)
        hint = f"💡 {td} стр. → *{up} ₽/стр*"
        if nt: hint += f"   |   от {nt} стр → *{np_} ₽/стр*"

        q.edit_message_text(
            f"📄 *Печать: {cn[color]}*\n\n"
            f"{hint}\n\n"
            f"{dot()}\n\n"
            f"Сколько *копий* напечатать?\n\n"
            f"_Напиши число или выбери кнопку:_",
            parse_mode="Markdown",
            reply_markup=kbd_qty("dq")
        )
        return SEL_DOC_QTY

    # ── количество кнопкой ──────────────────
    if data.startswith("pq_"):
        if not s: return do_cancel(uid, q, context)
        if data == "pq_hint":
            q.answer("Напиши число в чат — например: 3 или 10", show_alert=True)
            return SEL_PHOTO_QTY
        qty = int(data[3:])
        return _process_qty(uid, qty, "photo", context, q)

    if data.startswith("dq_"):
        if not s: return do_cancel(uid, q, context)
        if data == "dq_hint":
            q.answer("Напиши число в чат — например: 2 или 5", show_alert=True)
            return SEL_DOC_QTY
        qty = int(data[3:])
        return _process_qty(uid, qty, "doc", context, q)

    # ── доставка ────────────────────────────
    if data.startswith("dlv_"):
        if not s: return do_cancel(uid, q, context)
        method = data[4:]
        s["delivery_method"] = method
        dlabel, _ = DELIVERY_METHODS.get(method, ("Самовывоз",""))

        if method == "pickup":
            s["address"] = ""
            calc_total(uid)
            text = build_order_summary(uid)
            q.edit_message_text(text, reply_markup=kbd_confirm(), parse_mode="Markdown")
            return CONFIRMING
        else:
            c = get_client(uid)
            last_addr = c.get("last_address","") if c else ""
            rows = []
            if last_addr:
                rows.append([InlineKeyboardButton(
                    f"✓  Использовать: {last_addr[:35]}",
                    callback_data="use_last_addr")])
            rows.append([InlineKeyboardButton("✖  Отменить", callback_data="cancel")])
            q.edit_message_text(
                f"🚚 *{dlabel}*\n\n"
                f"📍 Напиши адрес доставки в чат"
                + (f"\n\n_Или используй последний адрес_" if last_addr else ""),
                parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup(rows)
            )
            return ENTER_ADDRESS

    if data == "use_last_addr":
        if not s: return do_cancel(uid, q, context)
        c = get_client(uid)
        s["address"] = c.get("last_address","") if c else ""
        calc_total(uid)
        text = build_order_summary(uid)
        q.edit_message_text(text, reply_markup=kbd_confirm(), parse_mode="Markdown")
        return CONFIRMING

    # ── подтверждение ───────────────────────
    if data == "confirm":
        if not s: return do_cancel(uid, q, context)
        ok, oid = save_order(uid, s)
        if ok:
            notify_admin(s, oid)
            ph = s["photos"]; dc = s["docs"]
            tp = s["total_photos"]; td = s["total_pages"]
            fn = {"small":"10×15/A6","medium":"13×18","large":"A4"}
            cn = {"bw":"Ч/Б","color":"Цветная"}
            dlabel = DELIVERY_METHODS.get(s["delivery_method"],("Самовывоз",""))[0]
            disc   = s.get("discount", 0)

            lines = [
                "🎊🎉 *Заказ оформлен!* 🎉🎊", "",
                f"🆔 `{oid}`",
                f"👤 {s['user_info']['first_name']}",
                "",
            ]
            if ph:
                lines.append(f"📸 Фото {fn.get(s['photo_format'],'')}  —  {tp * s['photo_qty']} шт.")
            if dc:
                lines.append(f"📄 Документы {cn.get(s['doc_color'],'')}  —  {td * s['doc_qty']} стр.")
            if disc:
                lines.append(f"👑 Скидка: {disc}%")
            lines += [
                "",
                f"💰 *Итого: {s['total']} ₽*",
                f"⏱ Срок: {estimate_delivery(tp+td)}",
                f"🚚 {dlabel}",
            ]
            if s.get("address"): lines.append(f"📍 {s['address']}")
            lines += [
                "",
                dot(),
                "",
                f"📞 {CONTACT_PHONE}",
                f"📌 Статус: {get_status_display('new')}",
                "🔔 Уведомлю при изменении статуса.",
                "",
                "💙 Спасибо за заказ!",
            ]
            # подсказка до следующей скидки
            c2 = get_client(uid)
            n2 = len(c2.get("orders",[])) if c2 else 0
            for thr, pct in sorted(LOYALTY_DISCOUNTS.items()):
                if n2 < thr:
                    lines.append(f"\n💡 Ещё *{thr-n2}* заказ(ов) — и скидка *{pct}%*")
                    break

            context.bot.send_message(chat_id=uid, text="\n".join(lines),
                                     parse_mode="Markdown")
            # превью фото
            if ph:
                try:
                    grp = []
                    for i, pf in enumerate(ph[:5]):
                        with open(pf["path"],"rb") as fp: raw = fp.read()
                        grp.append(InputMediaPhoto(raw,
                            caption=f"📸 Ваши фото ({len(ph)} шт.)" if i==0 else None))
                    if grp: context.bot.send_media_group(chat_id=uid, media=grp)
                except Exception as e:
                    logger.error(f"preview: {e}")
        else:
            context.bot.send_message(
                chat_id=uid,
                text=f"😔 Что-то пошло не так. Напишите нам: {CONTACT_PHONE}",
                parse_mode="Markdown"
            )

        cleanup(uid)
        try: q.message.delete()
        except: pass
        context.bot.send_message(chat_id=uid,
            text="Хочешь напечатать ещё что-нибудь?",
            reply_markup=kbd_after_order(uid))
        return COLLECTING_FILES

    return COLLECTING_FILES

# ═══════════════════════════════════════════
#  ВЕБ-ИНТЕРФЕЙС
# ═══════════════════════════════════════════
app = Flask(__name__)

CSS = """<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',Arial,sans-serif;background:#f0f2f5;min-height:100vh;padding:24px}
.wrap{max-width:1400px;margin:0 auto}
.topbar{background:#1a1a2e;border-radius:16px;padding:24px 32px;margin-bottom:24px;
        color:#fff;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:12px}
.topbar h1{font-size:1.6em;font-weight:700}
.nav{display:flex;gap:10px;flex-wrap:wrap}
.nav a{background:rgba(255,255,255,.12);color:#fff;text-decoration:none;
       padding:8px 20px;border-radius:10px;font-size:.9em;transition:background .2s}
.nav a:hover{background:rgba(255,255,255,.25)}
.stats-row{display:grid;grid-template-columns:repeat(auto-fill,minmax(160px,1fr));gap:16px;margin-bottom:24px}
.stat-card{background:#fff;border-radius:14px;padding:20px;text-align:center;
           box-shadow:0 2px 8px rgba(0,0,0,.06)}
.stat-card .val{font-size:2em;font-weight:700;color:#1a1a2e}
.stat-card .lbl{font-size:.8em;color:#888;margin-top:4px}
.grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(380px,1fr));gap:20px}
.card{background:#fff;border-radius:16px;overflow:hidden;box-shadow:0 2px 12px rgba(0,0,0,.07);
      transition:transform .15s,box-shadow .15s}
.card:hover{transform:translateY(-2px);box-shadow:0 6px 24px rgba(0,0,0,.12)}
.card-head{padding:16px 20px;background:#1a1a2e;color:#fff}
.card-head h2{font-size:.95em;word-break:break-all;margin-bottom:3px}
.card-head .sub{font-size:.78em;opacity:.7}
.card-body{padding:16px 20px}
.sbtn-row{display:flex;flex-wrap:wrap;gap:5px;margin-bottom:12px}
.sbtn{border:none;border-radius:8px;padding:4px 10px;font-size:.78em;
      font-weight:600;cursor:pointer;transition:opacity .15s}
.sbtn:hover{opacity:.8}
.s-new{background:#e0f0ff;color:#0066cc}
.s-processing{background:#fff3cd;color:#856404}
.s-printing{background:#d4edda;color:#155724}
.s-ready{background:#e8d5fb;color:#5a1e8c}
.s-shipped{background:#fce4f3;color:#8b1a5e}
.s-delivered{background:#c3f5e0;color:#0a5c38}
.s-cancelled{background:#fde8e8;color:#8b1a1a}
.gallery{display:flex;gap:8px;overflow-x:auto;padding:6px 0}
.thumb{width:70px;height:70px;object-fit:cover;border-radius:8px;cursor:pointer;
       flex-shrink:0;transition:transform .15s}
.thumb:hover{transform:scale(1.1)}
.tags{display:flex;gap:6px;flex-wrap:wrap;margin-top:8px}
.tag{background:#f0f2f5;border-radius:6px;padding:3px 10px;font-size:.78em;color:#555}
.acts{margin-top:12px;display:flex;gap:8px}
.btn{display:inline-block;padding:7px 16px;border-radius:9px;font-size:.85em;
     font-weight:600;text-decoration:none;color:#fff}
.btn-blue{background:#1a1a2e}
.btn-green{background:#198754}
table{width:100%;border-collapse:collapse;background:#fff;border-radius:16px;overflow:hidden;
      box-shadow:0 2px 8px rgba(0,0,0,.06)}
th{background:#1a1a2e;color:#fff;padding:12px 16px;text-align:left;font-weight:600}
td{padding:10px 16px;border-bottom:1px solid #f0f2f5;font-size:.9em}
tr:last-child td{border-bottom:none}
tr:hover td{background:#f8f9fa}
.badge{display:inline-block;padding:2px 10px;border-radius:20px;font-size:.78em;font-weight:600}
.badge-gold{background:#fff3cd;color:#856404}
pre{background:#f8f9fa;border-radius:10px;padding:14px;font-size:.82em;
    overflow-x:auto;white-space:pre-wrap;line-height:1.6;color:#333}
</style>"""

JS_STATUS = """<script>
function upd(id,s){
  fetch('/orders/'+id+'/status',{method:'POST',
    headers:{'Content-Type':'application/json'},
    body:JSON.stringify({status:s})
  }).then(r=>r.json()).then(d=>{if(d.success)location.reload();else alert('Ошибка');});
}
</script>"""

def status_btns(oid):
    def b(key, label, cls):
        return '<button class="sbtn ' + cls + '" onclick="upd(\"' + oid + '\",\"' + key + '\")">' + label + '</button>'
    parts = [
        '<div class="sbtn-row">',
        b("new","🆕 Новый","s-new"),
        b("processing","🔄 Обработка","s-processing"),
        b("printing","🖨 Печать","s-printing"),
        b("ready","✅ Готов","s-ready"),
        b("shipped","📦 Отправлен","s-shipped"),
        b("delivered","🏁 Доставлен","s-delivered"),
        b("cancelled","❌ Отменён","s-cancelled"),
        '</div>',
    ]
    return "".join(parts)

@app.route("/orders/")
def list_orders():
    try:
        orders = []
        if os.path.exists(ORDERS_PATH):
            hmap = {h.get("order_id"): h for h in load_history()}
            for item in sorted(os.listdir(ORDERS_PATH), reverse=True):
                ipath = os.path.join(ORDERS_PATH, item)
                if not os.path.isdir(ipath): continue
                files=[]; photos=[]; tsize=0
                for fn in os.listdir(ipath):
                    if fn == "информация_о_заказе.txt": continue
                    fp = os.path.join(ipath, fn); fsz = os.path.getsize(fp); tsize += fsz
                    iph = fn.lower().endswith((".jpg",".jpeg",".png"))
                    fi = {"name":fn,"size":format_size(fsz),"url":"/orders/"+item+"/"+fn,"is_photo":iph}
                    files.append(fi)
                    if iph: photos.append(fi)
                ctime = datetime.fromtimestamp(os.path.getctime(ipath))
                h = hmap.get(item, {})
                dlabel = DELIVERY_METHODS.get(h.get("delivery_method","pickup"), ("Самовывоз",""))[0]
                orders.append({
                    "id":item, "photos":photos[:4], "fc":len(files),
                    "size":format_size(tsize), "created":ctime.strftime("%d.%m.%Y %H:%M"),
                    "status":get_status_display(h.get("status","new")),
                    "total":h.get("total_price",""),
                    "delivery":dlabel,
                    "address":h.get("address",""),
                })
        orders.sort(key=lambda x:x["created"], reverse=True)
        cnt = len(orders)
        clients = load_clients()
        revenue = sum(o.get("total_price",0) for o in load_history() if isinstance(o.get("total_price"),int))

        # Генерируем карточки без вложенных f-строк
        cards_html = []
        for o in orders:
            # gallery
            gallery_items = []
            for p in o["photos"]:
                gallery_items.append(
                    '<img src="' + p["url"] + '" class="thumb" onclick="window.open('' + p["url"] + '')">'
                )
            gallery_html = '<div class="gallery">' + "".join(gallery_items) + '</div>' if o["photos"] else ""

            # tags
            tags = ['<span class="tag">📁 ' + str(o["fc"]) + ' файлов</span>',
                    '<span class="tag">💾 ' + o["size"] + '</span>']
            if o["total"]:
                tags.append('<span class="tag">💰 ' + str(o["total"]) + ' ₽</span>')
            tags.append('<span class="tag">🚚 ' + o["delivery"] + '</span>')
            tags_html = '<div class="tags">' + "".join(tags) + '</div>'

            addr_html = '<div style="font-size:.8em;color:#888;margin-top:6px">📍 ' + o["address"] + '</div>' if o["address"] else ""

            card = (
                '<div class="card">'
                '<div class="card-head">'
                '<h2>' + o["id"] + '</h2>'
                '<div class="sub">' + o["status"] + ' &bull; ' + o["created"] + '</div>'
                '</div>'
                '<div class="card-body">'
                + status_btns(o["id"]) +
                gallery_html +
                tags_html +
                addr_html +
                '<div class="acts">'
                '<a href="/orders/' + o["id"] + '/" class="btn btn-blue">Подробнее</a>'
                '<a href="/orders/' + o["id"] + '/download" class="btn btn-green">ZIP</a>'
                '</div>'
                '</div>'
                '</div>'
            )
            cards_html.append(card)

        html = (
            "<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
            "<title>Заказы — Print Bot</title>" + CSS + JS_STATUS + "</head><body>"
            '<div class="wrap">'
            '<div class="topbar">'
            "<h1>🖨 Print Bot — Заказы</h1>"
            '<div class="nav">'
            '<a href="/">Главная</a>'
            '<a href="/clients">Клиенты</a>'
            '<a href="/stats">API</a>'
            "</div></div>"
            '<div class="stats-row">'
            '<div class="stat-card"><div class="val">' + str(cnt) + '</div><div class="lbl">заказов</div></div>'
            '<div class="stat-card"><div class="val">' + str(len(clients)) + '</div><div class="lbl">клиентов</div></div>'
            '<div class="stat-card"><div class="val">' + str(revenue) + ' ₽</div><div class="lbl">выручка</div></div>'
            "</div>"
            '<div class="grid">' + "".join(cards_html) + "</div>"
            "</div></body></html>"
        )
        return html
    except Exception as e:
        return "Ошибка: " + str(e), 500

@app.route("/orders/<path:oid>/")
def view_order(oid):
    try:
        opath=os.path.join(ORDERS_PATH,oid)
        if not os.path.exists(opath): abort(404)
        info=""
        inf=os.path.join(opath,"информация_о_заказе.txt")
        if os.path.exists(inf):
            with open(inf,"r",encoding="utf-8") as f: info=f.read()
        status="new"
        for h in load_history():
            if h.get("order_id")==oid: status=h.get("status","new"); break
        files=[]; photos=[]; tsize=0
        for fn in sorted(os.listdir(opath)):
            if fn=="информация_о_заказе.txt": continue
            fp=os.path.join(opath,fn); fsz=os.path.getsize(fp); tsize+=fsz
            iph=fn.lower().endswith((".jpg",".jpeg",".png"))
            fi={"name":fn,"size":format_size(fsz),"url":f"/orders/{oid}/{fn}","is_photo":iph}
            files.append(fi); 
            if iph: photos.append(fi)
        ctime=datetime.fromtimestamp(os.path.getctime(opath))
        ph_html = "".join(
            '<img src="' + p["url"] + '" class="thumb" onclick="window.open(\"' + p["url"] + '\")" style="width:110px;height:110px">'
            for p in photos)
        fl_html = "".join(
            '<a href="' + f["url"] + '" download style="display:flex;align-items:center;gap:8px;padding:10px 14px;background:#f8f9fa;border-radius:10px;text-decoration:none;color:#333;margin-bottom:8px">'
            '<span style="font-size:1.4em">' + ("📸" if f["is_photo"] else "📄") + '</span>'
            '<div><div style="font-weight:600;font-size:.88em">' + f["name"] + '</div>'
            '<div style="font-size:.75em;color:#888">' + f["size"] + '</div></div></a>'
            for f in files)
        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
        <title>{oid}</title>{CSS}
        <script>function upd(s){{fetch('/orders/{oid}/status',{{method:'POST',
          headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{status:s}})
        }}).then(r=>r.json()).then(d=>{{if(d.success)location.reload();else alert('Ошибка');}});}}</script>
        </head><body><div class="wrap">
          <div class="topbar">
            <h1>📁 {oid}</h1>
            <div class="nav"><a href="/orders/">← Заказы</a><a href="/">Главная</a></div>
          </div>
          <div style="background:#fff;border-radius:16px;padding:24px;box-shadow:0 2px 8px rgba(0,0,0,.06)">
            <h3 style="margin-bottom:12px">Статус: {get_status_display(status)}</h3>
            {status_btns(oid)}
            <hr style="margin:20px 0;border:none;border-top:1px solid #f0f2f5">
            <h3 style="margin-bottom:12px">Информация</h3>
            <pre>{info}</pre>
            <hr style="margin:20px 0;border:none;border-top:1px solid #f0f2f5">
            <h3 style="margin-bottom:12px">Фото ({len(photos)})</h3>
            <div class="gallery" style="flex-wrap:wrap">{ph_html}</div>
            <hr style="margin:20px 0;border:none;border-top:1px solid #f0f2f5">
            <h3 style="margin-bottom:12px">Файлы ({len(files)})</h3>
            {fl_html}
            <div style="margin-top:20px">
              <a href="/orders/{oid}/download" class="btn btn-green" style="font-size:1em;padding:12px 28px">⬇ Скачать всё (ZIP)</a>
            </div>
          </div>
        </div></body></html>"""
    except Exception as e:
        return f"Ошибка: {e}", 500

@app.route("/clients")
def list_clients():
    try:
        clients=load_clients()
        clist=sorted(clients.values(),key=lambda c:len(c.get("orders",[])),reverse=True)
        rows=""
        for c in clist:
            n=len(c.get("orders",[])); spent=c.get("total_spent",0)
            d=0
            for thr,pct in sorted(LOYALTY_DISCOUNTS.items()):
                if n>=thr: d=pct
            badge=f'<span class="badge badge-gold">👑 {d}%</span>' if d else "—"
            fname = c.get("first_name",""); uname = c.get("username","")
            last = c.get("last_seen","")[:10]
            rows += ("<tr><td>" + fname + " @" + uname + "</td><td>" + str(n) + "</td><td>" + str(spent) + " ₽</td><td>" + badge + "</td><td style='color:#888'>" + last + "</td></tr>")
        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
        <title>Клиенты</title>{CSS}</head><body><div class="wrap">
          <div class="topbar"><h1>👥 Клиентская база</h1>
            <div class="nav"><a href="/orders/">Заказы</a><a href="/">Главная</a></div>
          </div>
          <div class="stats-row">
            <div class="stat-card"><div class="val">{len(clist)}</div><div class="lbl">клиентов</div></div>
          </div>
          <table><tr><th>Клиент</th><th>Заказов</th><th>Потрачено</th><th>Скидка</th><th>Последний визит</th></tr>
          {rows}</table>
        </div></body></html>"""
    except Exception as e:
        return f"Ошибка: {e}", 500

@app.route("/orders/<path:oid>/status", methods=["POST"])
def set_status(oid):
    try:
        data=request.get_json()
        if not data or not data.get("status"): return jsonify({"success":False}),400
        return jsonify({"success":update_order_status(oid,data["status"])})
    except Exception as e:
        return jsonify({"success":False,"error":str(e)}),500

@app.route("/orders/<path:oid>/download")
def dl_zip(oid):
    try:
        opath=os.path.join(ORDERS_PATH,oid)
        if not os.path.exists(opath): return "Не найден",404
        tmp=tempfile.NamedTemporaryFile(delete=False,suffix=".zip")
        with zipfile.ZipFile(tmp.name,"w") as zf:
            for root,_,files in os.walk(opath):
                for fn in files:
                    fp=os.path.join(root,fn); zf.write(fp,os.path.relpath(fp,opath))
        return send_file(tmp.name,as_attachment=True,download_name=f"{oid}.zip")
    except Exception as e:
        return f"Ошибка: {e}",500

@app.route("/orders/<path:oid>/<filename>")
def dl_file(oid,filename):
    try:
        return send_from_directory(os.path.join(ORDERS_PATH,oid),filename,as_attachment=True)
    except Exception as e:
        return f"Ошибка: {e}",500

@app.route("/webhook",methods=["POST"])
def webhook():
    global dispatcher
    try:
        if dispatcher is None: return jsonify({"error":"not ready"}),503
        data=request.get_json()
        if data:
            upd_obj=telegram.Update.de_json(data,bot)
            dispatcher.process_update(upd_obj)
        return "OK",200
    except Exception as e:
        logger.error(f"webhook: {e}"); return jsonify({"error":str(e)}),500

@app.route("/health")
def health():
    return jsonify({"status":"ok","bot_ready":dispatcher is not None,"sessions":len(user_sessions)})

@app.route("/reinit")
def reinit():
    global bot,updater,dispatcher
    bot=updater=dispatcher=None
    threading.Thread(target=_init_bot,daemon=True).start()
    return jsonify({"status":"reinit started"})

@app.route("/stats")
def stats():
    clients=load_clients()
    cnt=len([d for d in os.listdir(ORDERS_PATH) if os.path.isdir(os.path.join(ORDERS_PATH,d))]) if os.path.exists(ORDERS_PATH) else 0
    rev=sum(c.get("total_spent",0) for c in clients.values())
    return jsonify({"orders":cnt,"clients":len(clients),"revenue":rev,"sessions":len(user_sessions)})

@app.route("/")
def home():
    clients=load_clients(); now=datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    cnt=len([d for d in os.listdir(ORDERS_PATH) if os.path.isdir(os.path.join(ORDERS_PATH,d))]) if os.path.exists(ORDERS_PATH) else 0
    rev=sum(c.get("total_spent",0) for c in clients.values())
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
    <title>Print Bot</title>{CSS}</head><body>
    <div class="wrap" style="display:flex;align-items:center;justify-content:center;min-height:90vh">
      <div style="max-width:600px;width:100%">
        <div class="topbar" style="text-align:center;display:block;padding:40px">
          <h1 style="font-size:2.4em;margin-bottom:8px">🖨 Print Bot</h1>
          <p style="opacity:.7">Сервис быстрой печати фото и документов</p>
        </div>
        <div class="stats-row" style="grid-template-columns:repeat(3,1fr)">
          <div class="stat-card"><div class="val">{cnt}</div><div class="lbl">заказов</div></div>
          <div class="stat-card"><div class="val">{len(clients)}</div><div class="lbl">клиентов</div></div>
          <div class="stat-card"><div class="val">{rev} ₽</div><div class="lbl">выручка</div></div>
        </div>
        <div style="background:#fff;border-radius:16px;padding:24px;box-shadow:0 2px 8px rgba(0,0,0,.06)">
          <div class="nav" style="justify-content:center;margin-bottom:20px">
            <a href="/orders/" class="btn btn-blue" style="font-size:1em;padding:12px 28px">📦 Заказы</a>
            <a href="/clients" class="btn btn-blue" style="font-size:1em;padding:12px 28px">👥 Клиенты</a>
          </div>
          <div style="text-align:center;color:#888;font-size:.85em">
            📞 {CONTACT_PHONE}<br>⏰ {now}
          </div>
        </div>
      </div>
    </div></body></html>"""

# ═══════════════════════════════════════════
#  ИНИЦИАЛИЗАЦИЯ
# ═══════════════════════════════════════════
print("=" * 55)
print("🚀  Print Bot v6.0")
print(f"📁  {ORDERS_PATH}")
print(f"🤖  AI: {'✅' if OPENROUTER_API_KEY else '❌ нет OPENROUTER_API_KEY'}")
print("=" * 55)

def _build_conv():
    return ConversationHandler(
        entry_points=[
            CommandHandler("start",     cmd_start),
            CommandHandler("myorders",  cmd_myorders),
            CommandHandler("prices",    cmd_prices),
            CommandHandler("ai",        cmd_ai),
            CommandHandler("help",      cmd_help),
            MessageHandler(Filters.document | Filters.photo, handle_file),
        ],
        states={
            COLLECTING_FILES: [
                MessageHandler(Filters.document | Filters.photo, handle_file),
                CallbackQueryHandler(button_handler),
            ],
            SEL_PHOTO_FORMAT: [
                CallbackQueryHandler(button_handler, pattern="^(pf_|add_more|cancel)"),
            ],
            SEL_PHOTO_QTY: [
                MessageHandler(Filters.text & ~Filters.command, qty_text_handler),
                CallbackQueryHandler(button_handler, pattern="^(pq_|cancel)"),
            ],
            SEL_DOC_TYPE: [
                CallbackQueryHandler(button_handler, pattern="^(dt_|add_more|cancel)"),
            ],
            SEL_DOC_QTY: [
                MessageHandler(Filters.text & ~Filters.command, qty_text_handler),
                CallbackQueryHandler(button_handler, pattern="^(dq_|cancel)"),
            ],
            SEL_DELIVERY: [
                CallbackQueryHandler(button_handler, pattern="^(dlv_|cancel)"),
            ],
            ENTER_ADDRESS: [
                MessageHandler(Filters.text & ~Filters.command, address_handler),
                CallbackQueryHandler(button_handler, pattern="^(use_last_addr|cancel)"),
            ],
            CONFIRMING: [
                CallbackQueryHandler(button_handler, pattern="^(confirm|cancel)"),
            ],
            AI_CHAT: [
                MessageHandler(Filters.text & ~Filters.command, ai_msg_handler),
                CallbackQueryHandler(button_handler),
            ],
        },
        fallbacks=[
            CommandHandler("start",    cmd_start),
            CommandHandler("myorders", cmd_myorders),
            CommandHandler("prices",   cmd_prices),
            CommandHandler("ai",       cmd_ai),
            CommandHandler("help",     cmd_help),
        ],
        allow_reentry=True,
    )

def _init_bot():
    global bot, updater, dispatcher
    if dispatcher is not None: return
    for attempt in range(1, 11):
        try:
            print(f"🤖  Инициализация (попытка {attempt}/10)...")
            bot        = telegram.Bot(token=TOKEN)
            updater    = Updater(token=TOKEN, use_context=True,
                                 request_kwargs={"connect_timeout":30,"read_timeout":30})
            dispatcher = updater.dispatcher
            dispatcher.add_handler(_build_conv())
            bot.set_webhook(url=f"{RENDER_URL}/webhook")
            print(f"✅  Webhook: {RENDER_URL}/webhook")
            print("✅  БОТ ГОТОВ!")
            return
        except Exception as e:
            print(f"❌  Попытка {attempt}: {str(e)[:100]}")
            bot = updater = dispatcher = None
            if attempt < 10: _time.sleep(5 * attempt)
    print("❌  Не удалось запустить бота.")

_init_bot()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT)
